import time

from docx import Document
from docx.shared import Inches

months = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября", "ноября",
          "декабря"]


def create_dox(f,
               i,
               o,
               group):
    # Создание нового документа
    doc = Document()

    # Открытие шаблона
    template = Document('шаблон.docx')

    # Копирование содержимого шаблона в новый документ
    for element in template.element.body:
        doc.element.body.append(element)

    # f = input("Введите фамилию: ")
    # i = input("Введите имя: ")
    # o = input("Введите отчество: ")
    # group = input("Введите группу: ")
    day = int(time.strftime(f'%d'))
    month = int(time.strftime("%m"))

    # Замена текста в шаблоне
    table = doc.tables[0]

    for row in table.rows:
        for cell in row.cells:
            if 'ФИО' in cell.text:
                cell.text = cell.text.replace('ФИО', f"{f} {i[0]}.{o[0]}.")
            if 'группа' in cell.text:
                cell.text = cell.text.replace('группа', group)
            if 'день' in cell.text:
                cell.text = cell.text.replace('день', str(day))
            if 'месяц' in cell.text:
                cell.text = cell.text.replace('месяц', months[month - 1])

    word_to_replace = 'ФОТО'

    # Путь к новому изображению
    count = 1
    for paragraph in doc.paragraphs:
        # Поиск слова для замены в каждом параграфе
        if word_to_replace in paragraph.text:
            # Создание нового рисунка и добавление его вместо слова

            # Удаление слова из параграфа
            paragraph.text = paragraph.text.replace(word_to_replace, '')
            run = paragraph.add_run()
            run.add_picture(f"{count}.png", width=Inches(5.0), height=Inches(3.0))
            count += 1

    # Сохранение нового документа
    doc.save(f'ЦТвПД_{group}_{f}_{i[0]}{o[0]}_ТЕСТ.doc')
