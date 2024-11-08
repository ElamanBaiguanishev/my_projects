import json
import time
import pyautogui as pg
from docx import Document
from docx.shared import Inches


with open('data_ct.json', 'r', encoding='utf-8') as file:
    file_content = file.read()

# Преобразуем содержимое файла в словарь
data_ct = json.loads(file_content)

print(data_ct)

months = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября", "ноября",
          "декабря"]


def create_dox(f,
               i,
               o,
               group):
    # Создание нового документа
    doc = Document()

    # Открытие шаблона
    template = Document('шаблон.docx')

    # Копирование содержимого шаблона в новый документ
    for element in template.element.body:
        doc.element.body.append(element)

    # f = input("Введите фамилию: ")
    # i = input("Введите имя: ")
    # o = input("Введите отчество: ")
    # group = input("Введите группу: ")
    day = int(time.strftime(f'%d'))
    month = int(time.strftime("%m"))

    # Замена текста в шаблоне
    table = doc.tables[0]

    for row in table.rows:
        for cell in row.cells:
            if 'ФИО' in cell.text:
                cell.text = cell.text.replace('ФИО', f"{f} {i[0]}.{o[0]}.")
            if 'группа' in cell.text:
                cell.text = cell.text.replace('группа', group)
            if 'день' in cell.text:
                cell.text = cell.text.replace('день', str(day))
            if 'месяц' in cell.text:
                cell.text = cell.text.replace('месяц', months[month - 1])

    word_to_replace = 'ФОТО'

    # Путь к новому изображению
    count = 1
    for paragraph in doc.paragraphs:
        # Поиск слова для замены в каждом параграфе
        if word_to_replace in paragraph.text:
            # Создание нового рисунка и добавление его вместо слова

            # Удаление слова из параграфа
            paragraph.text = paragraph.text.replace(word_to_replace, '')
            run = paragraph.add_run()
            run.add_picture(f"{count}.png", width=Inches(5.0), height=Inches(3.0))
            count = count + 1

    # Сохранение нового документа
    doc.save(f'D:\\work\\Цифровые технологии\\тест_ЦТ\\Бот\\ЦТвПД_{group}_{f}_{i[0]}{o[0]}_ТЕСТ.doc')


layout = dict(zip(map(ord, "йцукенгшщзхъфывапролджэячсмитьбю.ё"
                           'ЙЦУКЕНГШЩЗХЪФЫВАПРОЛДЖЭЯЧСМИТЬБЮ,Ё'),
                  "qwertyuiop[]asdfghjkl;'zxcvbnm,./`"
                  'QWERTYUIOP{}ASDFGHJKL:"ZXCVBNM<>?~'))

screen_width, screen_height = pg.size()

center_x = screen_width // 2
center_y = screen_height // 2

second_name = data_ct["f"]
first_name = data_ct["i"]
o = data_ct["o"]
group = data_ct["group"]
adapter = data_ct["adapter"]

pg.hotkey("win", "s")
time.sleep(1)
pg.hotkey("shift", "alt")

pg.typewrite('Семья'.translate(layout), interval=0.1)
pg.press('enter')
time.sleep(2)
pg.moveTo(x=508, y=641)
pg.doubleClick()
pg.moveTo(800, 650)

time.sleep(5)
pg.doubleClick()
pg.moveTo(800, 610)
time.sleep(5)
pg.doubleClick()
time.sleep(5)
pg.typewrite(second_name.translate(layout))
pg.moveTo(x=920, y=110)
pg.mouseDown()
time.sleep(1)
pg.dragTo(x=709, y=161, duration=1)
pg.mouseUp()
pg.screenshot("1.png")
pg.hotkey("shift", "alt")
time.sleep(1)

pg.hotkey("win", "s")
time.sleep(1)
pg.typewrite('cmd', interval=0.1)
time.sleep(1)
pg.press('enter')
time.sleep(2)


pg.doubleClick(center_x, center_y)
# pg.click()
pg.hotkey("win", "up")
pg.typewrite('dir.exe', interval=0.2)
pg.press('enter')
pg.hotkey("shift", "alt")
pg.typewrite(f'{second_name}'.translate(layout), interval=0.2)
pg.press('enter')
pg.typewrite(f'{first_name}'.translate(layout), interval=0.2)
pg.press('enter')

pg.hotkey("shift", "alt")
pg.typewrite(adapter, interval=0.2)
pg.press('enter')
pg.scroll(1400)
time.sleep(1)
pg.screenshot("2.png")
pg.scroll(-300)
time.sleep(1)
pg.screenshot("3.png")
pg.scroll(-1100)
time.sleep(1)
pg.screenshot("4.png")

create_dox(second_name, first_name, o, group)
